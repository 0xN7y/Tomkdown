#N7y
import argparse
from pathlib import Path
import fitz  
from docx import Document
import sys
if len(sys.argv) == 1:
    print("""
            \r \r  Convert data from PDF, Docs into an Markdown
\t\t\t\t\t Auther: N7y

[...     [..                   
[. [..   [..[..... [..         
[.. [..  [..      [.. [..   [..
[..  [.. [..     [..   [.. [.. 
[..   [. [..    [..      [...  
[..    [. ..    [..       [..  
[..      [..    [..      [..   
                       [..  

    \r Usage is very human just
    \r python3 tomarkdown.py --help # for more info

        """)
    exit()

def ex_pdfimg(pdf_path, output_dir):
    images = []
    doc = fitz.open(pdf_path)
    for page_number in range(len(doc)):
        page = doc[page_number]
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_filename = f"{output_dir}/page{page_number+1}_img{img_index+1}.{image_ext}"
            with open(image_filename, "wb") as image_file:
                image_file.write(image_bytes)
            images.append(image_filename)
    return images

def ptomd(pdf_path, output_dir):
    doc = fitz.open(pdf_path)
    md_content = []

    images = ex_pdfimg(pdf_path, output_dir)

    for page_number in range(len(doc)):
        page = doc[page_number]
        text = page.get_text("dict")
        blocks = text.get("blocks", [])
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        if span["flags"] & 2:  # Bold text
                            md_content.append(f"**{span['text']}**")
                        elif span["flags"] & 4:  # Italic text
                            md_content.append(f"*{span['text']}*")
                        else:
                            md_content.append(span["text"])
                        md_content.append(" ")
        md_content.append("\n\n")

    
    for img in images:
        md_content.append(f"![Image](./{Path(img).name})")

    return "\n".join(md_content)

def doctomd(docx_path):
    doc = Document(docx_path)
    md_content = []

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            level = int(paragraph.style.name.replace("Heading ", ""))
            md_content.append(f"{'#' * level} {paragraph.text}")
        elif any(run.bold for run in paragraph.runs):
            md_content.append(f"**{paragraph.text}**")
        elif any(run.italic for run in paragraph.runs):
            md_content.append(f"*{paragraph.text}*")
        elif paragraph.style.name == "List Bullet":
            md_content.append(f"- {paragraph.text}")
        elif paragraph.style.name == "List Number":
            md_content.append(f"1. {paragraph.text}")
        else:
            md_content.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_content = " | ".join(cell.text.strip() for cell in row.cells)
            md_content.append(row_content)
            md_content.append("-" * len(row_content))

    return "\n".join(md_content)




parser = argparse.ArgumentParser(description="Convert PDF and DOCX files to Markdown.")
parser.add_argument("file_", help="Path to the pdf or docs File.")
parser.add_argument("output_dir", help="Directory to save the output Markdown file.")

args = parser.parse_args()

file_ = Path(args.file_)
# print(file_)
output_dir = Path(args.output_dir)
# print(output_dir)

# if dir not exist
output_dir.mkdir(parents=True, exist_ok=True)

if not file_.exists():
    print("file does not exist")
    exit()

if file_.suffix.lower() == ".pdf":
    md_content = ptomd(file_, output_dir)
elif file_.suffix.lower() == ".docx":
    md_content = doctomd(file_)
else:
    print("Only pdf or docs supported!...")
    exit()

output_file = f"{output_dir}/{file_.stem}.md"



with open(output_file, "w", encoding="utf-8") as f:
    f.write(md_content)

print(f"Generated Markdown  {output_file}")

